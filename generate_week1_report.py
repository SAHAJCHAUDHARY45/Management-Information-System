from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()

doc.add_heading('College Management Information System (MIS) using Django', 0)
doc.add_paragraph('Week 1 Progress Report')

doc.add_heading('Objective & Problem Statement', level=1)
doc.add_paragraph(
    'The objective of this project is to develop a web-based College Management Information System (MIS) using Django. '
    'The system will streamline student and faculty management, announcements, subject and result tracking, and provide dashboards for all users.\n'
    'Problem Statement: Manual management of student records, announcements, and results is time-consuming and error-prone. '
    'A digital MIS will improve efficiency and transparency.'
)

doc.add_heading('Project Scope & Features', level=1)
doc.add_paragraph('- Student registration and management\n'
                 '- Faculty registration and management\n'
                 '- Subject and result management\n'
                 '- Announcements (individual/group)\n'
                 '- Separate dashboards for students and faculty\n'
                 '- Admin panel for full control')

doc.add_heading('Technology Stack', level=1)
doc.add_paragraph('Python 3.x\nDjango 5.x\nBootstrap 5 (for frontend)\nSQLite (default Django DB)')

doc.add_heading('Timeline', level=1)
table = doc.add_table(rows=1, cols=2)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Week'
hdr_cells[1].text = 'Tasks'
tasks = [
    ('1', 'Project setup, models, admin, basic pages'),
    ('2', 'Registration, login, dashboards'),
    ('3', 'Announcements, subject/result management'),
    ('4', 'Final polish, testing, documentation'),
]
for week, task in tasks:
    row_cells = table.add_row().cells
    row_cells[0].text = week
    row_cells[1].text = task

doc.add_heading('Initial Setup', level=1)
doc.add_paragraph('- Django project created: college_mis\n'
                 '- Apps created: core, students, faculty, announcements\n'
                 '- Custom user model: User (with is_student and is_faculty)\n'
                 '- Basic models: Student, Faculty, Subject, Result, Announcement\n'
                 '- Admin panel: All models registered and manageable\n'
                 '- Basic navigation: Home, login, registration, dashboard URLs set up')

doc.add_heading('Code Snippets', level=1)
doc.add_paragraph('models.py (core):', style='Intense Quote')
doc.add_paragraph('''class User(AbstractUser):\n    is_student = models.BooleanField(default=False)\n    is_faculty = models.BooleanField(default=False)\n\nclass Student(models.Model):\n    user = models.OneToOneField('core.User', on_delete=models.CASCADE)\n    roll_number = models.CharField(max_length=20, unique=True)\n    phone_number = models.CharField(max_length=20, blank=True)\n    # ... other fields ...''')

doc.add_paragraph('settings.py:', style='Intense Quote')
doc.add_paragraph('''INSTALLED_APPS = [\n    # ... default apps ...\n    'core',\n    'students',\n    'faculty',\n    'announcements',\n]\nAUTH_USER_MODEL = 'core.User'\n''')

doc.add_paragraph('urls.py:', style='Intense Quote')
doc.add_paragraph('''urlpatterns = [\n    path('', home, name='home'),\n    path('admin/', admin.site.urls),\n    path('students/', include('students.urls')),\n    path('faculty/', include('faculty.urls')),\n    path('announcements/', include('announcements.urls')),\n    path('login/', auth_views.LoginView.as_view(template_name='login.html'), name='login'),\n    path('logout/', auth_views.LogoutView.as_view(next_page='login'), name='logout'),\n]\n''')

doc.add_heading('Next Steps', level=1)
doc.add_paragraph('- Implement registration forms and dashboards\n'
                 '- Add subject and result management\n'
                 '- Add announcement features')

doc.save('MIS_Week1_Report.docx')
print('Report generated: MIS_Week1_Report.docx') 